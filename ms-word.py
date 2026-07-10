#!/usr/bin/env python3
# -*- coding: utf-8 -*-
r"""
msword_mcp.py - A single-file MCP (Model Context Protocol) stdio server that
gives an AI agent read/search/edit/generate access to Word .docx files.

It follows a simple open -> edit -> save workflow (msword_open ... msword_save),
with the .docx engine provided by the Python library `python-docx`.

WHAT IT CAN DO
    - Open a .docx into an in-memory session, keyed by session_id.
    - Read full content (linear text, or a structured block list).
    - Search text (body paragraphs and table cells).
    - Replace text, correctly handling matches that span multiple runs.
    - Record edits as Word TRACKED CHANGES (revisions) with author + date,
      the way Word itself records them:
        * replacements are DIFFED WORD-BY-WORD, so only the words that
          actually change are marked (never "whole paragraph deleted +
          whole paragraph inserted"); unchanged words keep their original
          runs and formatting
        * whole-paragraph insert/delete including the PARAGRAPH MARK, so
          accepting/rejecting adds or removes the paragraph itself
        * rewrite a whole paragraph (set_paragraph_text) as a tracked
          word-level diff of old vs new
    - List tracked changes (with location + context), then accept or reject
      them - all at once, or individually by change id.
    - Read/search the FINAL view while changes are pending (like Word's
      "No Markup"): pending insertions are visible, pending deletions hidden.
    - Append paragraphs, headings and tables; insert paragraphs anywhere.
    - Apply simple paragraph formatting (bold/italic/underline/alignment/style).
    - Save in place, or save-as to a new path.

TRACKED CHANGES - SCOPE AND VERIFICATION
    Supported (mirrors what Word records for normal typing with Track
    Changes on):
      - text insertions/deletions: <w:ins> / <w:del> + <w:delText>, produced
        from a word-level diff of old vs new text (difflib.SequenceMatcher
        over word/whitespace tokens);
      - paragraph insertion/deletion including the paragraph mark
        (an empty <w:ins>/<w:del> inside w:pPr/w:rPr), so a deleted
        paragraph disappears entirely on accept and an inserted paragraph
        disappears entirely on reject, with Word's paragraph-merge
        semantics (the following paragraph's mark survives a merge);
      - accept/reject of every change, or of individual changes by w:id.
    NOT supported, by design (fragile/rare OOXML cases): tracked moves
    (w:moveFrom/w:moveTo - a move appears as delete + insert instead),
    tracked formatting-only changes (w:rPrChange/w:pPrChange), tracked
    table row/cell revisions, and comment threads.

    Verified before delivery: correct <w:ins>/<w:del>/<w:delText> and
    paragraph-mark structure with unique ids + author + date; word-level
    diff granularity; cross-run matches with formatting preserved on
    unchanged words; accept-all / reject-all / accept-by-id round-trips
    (all exercised by --check); and an independent round-trip through
    LibreOffice Writer, which parsed and preserved the revisions and their
    authors. (One known LibreOffice-only quirk: if the SAME paragraph mark
    carries both an insertion and a deletion - e.g. a tracked-inserted
    paragraph later tracked-deleted - LO's re-export keeps only one of the
    two markers. The files this server writes are correct; only re-saving
    them from LO loses that corner case.) The ONE thing not verifiable off
    your network is the exact Microsoft Word review UI - open one test file
    in Word on your endpoint and confirm the changes appear under Review
    and that Accept/Reject behave as expected before relying on it for
    anything important.

WHAT IT CANNOT DO
    - Comment threads (no reliable python-docx API).
    - Tracked move / formatting-only / table-row revisions (see scope above).

=============================================================================
 DEPENDENCIES  (this is the ONE place this project leaves the standard library)
=============================================================================
    python-docx   (import name: docx)   -> the .docx engine
      +- lxml            (compiled C extension, pulled in by python-docx)
      +- typing_extensions

    AIRGAPPED WINDOWS INSTALL (sideload wheels, same pattern as pymupdf):
      1. On an internet-connected box, download wheels for your EXACT
         interpreter. lxml is compiled, so the wheel must match Python
         version + architecture, e.g. for CPython 3.11 64-bit:
             lxml-<ver>-cp311-cp311-win_amd64.whl
         python-docx and typing_extensions are pure-Python (any wheel works).
         A reliable way to grab the correct set in one go, run ON the target
         Python version if you can:
             python -m pip download python-docx -d .\wheels
      2. Transfer the .\wheels folder to the airgapped endpoint.
      3. Install with the SAME interpreter Continue will launch (use -m pip so
         the interpreter and pip cannot drift apart):
             "C:\path\to\python.exe" -m pip install --no-index ^
                 --find-links .\wheels python-docx
      4. Confirm the interpreter can see it:
             "C:\path\to\python.exe" -m docx  (no error = importable)

    If your corporate mirror proxies PyPI you may be able to skip the manual
    download and just run:  python -m pip install python-docx
    lxml wheel/interpreter mismatch is the #1 failure here; this server logs
    sys.executable and both versions at startup so a mismatch is obvious.

=============================================================================
 VALIDATE BEFORE WIRING IN  (single-transfer sanity check)
=============================================================================
    Run the built-in self-test. It creates a temp .docx, opens/edits/saves/
    reopens it, and prints PASS/FAIL. No arguments, no network, no side files
    left behind:
        "C:\path\to\python.exe" msword_mcp.py --check

    Expected tail of output on success:
        [check] round-trip: PASS
        [check] ALL CHECKS PASSED

=============================================================================
 WIRE INTO CONTINUE  (config.yaml)
=============================================================================
    Add under mcpServers. Use the SAME python.exe you installed the wheels
    with, and set PYTHONUTF8 so Windows cp1252 cannot corrupt the stdio JSON
    stream. After editing config.yaml, run "Developer: Reload Window" rather
    than toggling the server (avoids the "already connected to transport" bug).

        mcpServers:
          - name: msword-py
            command: C:\path\to\python.exe
            args:
              - C:\path\to\msword_mcp.py
              - --author
              - Matt
              - --document-root
              - C:\Users\me\Documents\ai_docs
            env:
              PYTHONUTF8: "1"

    The --author value is stamped on every tracked change. Omit the two --author
    lines to fall back to the TRACKED_CHANGE_AUTHOR config constant below.
    The --document-root folder is REQUIRED (here, via MSWORD_DOCUMENT_ROOT, or
    via the DOCUMENT_ROOT constant): all open/save paths must be inside it and
    the server refuses to start without one.

    (If your Continue build uses the older command/args-in-one style, match
    whatever your existing working Python MCP servers use - the launch shape
    is identical to them.)

=============================================================================
 PROTOCOL / TRANSPORT NOTES
=============================================================================
    - MCP stdio transport = newline-delimited JSON-RPC 2.0 on stdin/stdout.
    - stdout is SACRED: only JSON-RPC messages go there, one per line. Every
      diagnostic goes to stderr via log(). Any stray print() to stdout would
      corrupt the stream.
    - Streams are reconfigured to UTF-8 in-script as a belt-and-braces measure
      alongside PYTHONUTF8.

Author's assumptions (flagged per the airgap "a caveat is cheaper than a
failed transfer" rule):
    - A "reasonably modern" python-docx (>= 0.8.11; tested here on 1.2.0).
    - Editing scope for search/replace is the document BODY paragraphs and
      TABLE cells. Headers/footers are intentionally NOT edited (predictable,
      and they rarely carry the target content). Say the word if you need them.
    - When a replaced match spans runs with different formatting, the inserted
      text takes the formatting of the run where the match STARTS. This is the
      standard, unavoidable trade-off for run-aware replacement.
"""

# =============================================================================
# CONFIGURATION  (all user-editable settings live here, nothing scattered below)
# =============================================================================
SERVER_NAME = "msword-py"          # advertised to the MCP client
SERVER_VERSION = "1.3.0"
PROTOCOL_VERSION_FALLBACK = "2024-11-05"  # used if the client sends none

# REQUIRED path sandbox. The server refuses to open or save any file outside
# this directory tree, and REFUSES TO START if no root is configured - the
# model chooses open/save paths, so an unconfined server could read/write any
# .docx this account can. Set it here, or at launch with --document-root or
# the MSWORD_DOCUMENT_ROOT environment variable (which take priority over
# this constant). Symlinks are resolved before the containment check.
#   e.g. DOCUMENT_ROOT = r"C:\Users\you\Documents\ai_docs"
# (--check is exempt: the self-test sandboxes itself to its own temp folder.)
# Related caution: only open .docx files from trusted sources - a maliciously
# crafted file could use XML entity tricks to pull local file contents into
# the document text that the model then reads.
DOCUMENT_ROOT = None

MAX_SESSIONS = 32                    # guard against runaway open() calls
SEARCH_CONTEXT_CHARS = 40            # chars of context either side of a match

# Default author name stamped on tracked changes (w:author). Override per
# launch with:  --author "Matt"  in Continue's args: block. The date on each
# change is always the current date, computed at edit time.
TRACKED_CHANGE_AUTHOR = "AI Assistant (Continue)"
# =============================================================================

import sys
import os
import re
import json
import uuid
import copy
import difflib
import argparse
import traceback
import datetime

# --- Make stdio UTF-8 regardless of the Windows console codepage. -----------
# Belt-and-braces alongside PYTHONUTF8=1 in the Continue env: block.
for _stream in ("stdin", "stdout", "stderr"):
    try:
        getattr(sys, _stream).reconfigure(encoding="utf-8")
    except Exception:
        # Older/edge interpreters may not support reconfigure; PYTHONUTF8 covers it.
        pass


def log(msg):
    """All diagnostics go to stderr ONLY. stdout is reserved for JSON-RPC."""
    try:
        sys.stderr.write("[{}] {}\n".format(SERVER_NAME, msg))
        sys.stderr.flush()
    except Exception:
        pass


# --- Import the engine, failing loudly and specifically on mismatch. --------
try:
    import docx  # python-docx
    from docx.document import Document as _DocumentClass
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.exceptions import PackageNotFoundError
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception as _imp_err:  # pragma: no cover - exercised only on a broken install
    sys.stderr.write(
        "[{}] FATAL: could not import python-docx.\n"
        "        Interpreter : {}\n"
        "        Error       : {}\n"
        "        Fix: install python-docx into THIS interpreter "
        "(see the docstring's airgapped install steps).\n".format(
            SERVER_NAME, sys.executable, _imp_err
        )
    )
    sys.exit(1)


def _versions():
    """Best-effort version strings for startup diagnostics."""
    try:
        from importlib.metadata import version
        docx_ver = version("python-docx")
    except Exception:
        docx_ver = "unknown"
    try:
        from lxml import etree
        lxml_ver = etree.__version__
    except Exception:
        lxml_ver = "unknown"
    return docx_ver, lxml_ver


# =============================================================================
# SESSION STATE
# =============================================================================
# session_id -> {"path": str, "doc": docx Document, "opened_at": iso str}
SESSIONS = {}

# Author stamped on tracked changes. Seeded from config, optionally replaced
# by the --author launch flag in main().
AUTHOR = TRACKED_CHANGE_AUTHOR

# XML namespace literal for the reserved 'xml' prefix (qn() does not map it).
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


class ToolError(Exception):
    """Raised by a tool handler to return a clean isError result to the client."""


def _now_iso():
    return datetime.datetime.now().isoformat(timespec="seconds")


def _require(args, key):
    """Fetch a required argument or raise a clear ToolError."""
    if key not in args or args[key] is None:
        raise ToolError("Missing required argument: '{}'".format(key))
    return args[key]


def _resolve_path(path):
    """
    Expand + absolutise a path and enforce the DOCUMENT_ROOT sandbox.
    realpath (not just abspath) so a symlink inside the root cannot point the
    server at a file outside it.
    """
    if not isinstance(path, str) or not path.strip():
        raise ToolError("Path must be a non-empty string")
    rp = os.path.realpath(os.path.expanduser(path))
    if not DOCUMENT_ROOT:
        # main() refuses to start without a root; this guards direct callers.
        raise ToolError("No DOCUMENT_ROOT is configured; file access is disabled.")
    root = os.path.realpath(os.path.expanduser(DOCUMENT_ROOT))
    try:
        common = os.path.commonpath([root, rp])
    except ValueError:
        # Different drives on Windows raise ValueError from commonpath.
        common = None
    if common != root:
        raise ToolError(
            "Path is outside the permitted DOCUMENT_ROOT and was refused."
        )
    return rp


def _get_session(args):
    """Return the session dict for a required session_id argument."""
    sid = _require(args, "session_id")
    session = SESSIONS.get(sid)
    if session is None:
        raise ToolError(
            "Unknown session_id '{}'. Call msword_open first.".format(sid)
        )
    return session


# =============================================================================
# DOCUMENT HELPERS
# =============================================================================
def _iter_block_items(parent):
    """
    Yield Paragraph and Table objects in true document order.

    python-docx exposes document.paragraphs and document.tables as SEPARATE
    flat lists, losing their interleaving. This walks the underlying XML body
    so blocks come back in the order they actually appear.
    """
    if isinstance(parent, _DocumentClass):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ToolError("Unsupported parent for block iteration")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _all_editable_paragraphs(doc):
    """
    Every paragraph we are willing to search/replace within: body paragraphs
    plus paragraphs inside every table cell (recursively for nested tables).
    Order is not guaranteed to be strict document order, which is fine because
    replace only needs to visit each paragraph once.
    """
    def _cell_paragraphs(cell):
        for p in cell.paragraphs:
            yield p
        for tbl in cell.tables:  # nested tables
            for row in tbl.rows:
                for c in row.cells:
                    for p in _cell_paragraphs(c):
                        yield p

    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in _cell_paragraphs(cell):
                    yield p


def _replace_in_paragraph(paragraph, find, replace, remaining):
    """
    Replace occurrences of `find` with `replace` inside a single paragraph,
    correctly handling matches that span multiple runs.

    `remaining` caps how many replacements are still allowed (None = no cap).
    Returns the number of replacements made in this paragraph.

    Method: on each pass, concatenate all run texts, locate the first match in
    the combined text, then rewrite each run so the matched characters are
    removed from the runs they fell in and the replacement text is inserted
    into the run where the match STARTS. That start run's formatting is what
    the replacement inherits (the documented trade-off for cross-run matches).
    """
    made = 0
    while True:
        if remaining is not None and made >= remaining:
            break
        runs = paragraph.runs
        if not runs:
            break
        texts = [r.text for r in runs]
        full = "".join(texts)
        idx = full.find(find)
        if idx == -1:
            break
        end = idx + len(find)

        # Absolute start offset of each run within the combined text.
        starts = []
        acc = 0
        for t in texts:
            starts.append(acc)
            acc += len(t)

        for i, run in enumerate(runs):
            rstart = starts[i]
            rtext = texts[i]
            rend = rstart + len(rtext)
            # Kept slice to the LEFT of the match within this run.
            left = rtext[: max(0, min(rend, idx) - rstart)] if rstart < idx else ""
            # Kept slice to the RIGHT of the match within this run.
            right = rtext[max(0, end - rstart):] if rend > end else ""
            if rstart <= idx < rend:
                # This run contains the match start -> insertion happens here.
                run.text = left + replace + right
            else:
                run.text = left + right
        made += 1
    return made


# -----------------------------------------------------------------------------
# TRACKED CHANGES (revisions)
#
# Tracked changes are not a python-docx feature or a document flag - they are
# specific OOXML elements inside the paragraph:
#   text insertion : run(s) wrapped in  <w:ins w:id w:author w:date> ... </w:ins>
#   text deletion  : run(s) wrapped in  <w:del ...>, and each run's <w:t>
#                    becomes <w:delText> (text is stored but shown struck out)
#   paragraph mark : an empty <w:ins>/<w:del> inside w:pPr/w:rPr marks the
#                    pilcrow itself as inserted/deleted - this is how Word
#                    records whole-paragraph insertion/deletion, and it is what
#                    makes accept/reject add or remove the paragraph itself.
#
# To match how Word records an edit, a "replace" is NOT one big delete plus one
# big insert: the old and new text are diffed word-by-word and only the words
# that actually change are wrapped. Unchanged words keep their original runs
# (and therefore their formatting) untouched.
#
# We build these elements directly via lxml (OxmlElement/qn), which python-docx
# exposes on every element. Excluded by design: tracked moves, formatting-only
# revisions and table-row revisions (the fragile OOXML cases).
# -----------------------------------------------------------------------------
def _today_iso():
    """Current date as an OOXML xs:dateTime (time zeroed, UTC 'Z')."""
    return datetime.date.today().strftime("%Y-%m-%dT00:00:00Z")


# Tokeniser for the word-level diff: alternating words / whitespace stretches.
_TOKEN_RE = re.compile(r"\s+|\S+")


def _word_diff(old, new):
    """
    Word-level diff between two strings, expressed as edits to `old`.

    Returns (del_spans, insertions):
      del_spans  - sorted list of (start, end) char offsets into `old` that
                   must be recorded as tracked deletions
      insertions - dict {char_offset_into_old: text} of tracked insertions,
                   each anchored immediately BEFORE the character at that
                   offset (offset == len(old) means "append at the end").
                   Replacement text is anchored at the END of its deleted
                   span so the markup reads: old text struck out, then new.
    """
    a = _TOKEN_RE.findall(old)
    b = _TOKEN_RE.findall(new)
    # autojunk=False: whitespace tokens are "popular" and would otherwise be
    # junked, which stops the words around them from lining up as equal.
    matcher = difflib.SequenceMatcher(a=a, b=b, autojunk=False)
    # a_off[i] = char offset of token i within `old`.
    a_off = [0]
    for tok in a:
        a_off.append(a_off[-1] + len(tok))

    del_spans = []
    insertions = {}

    def _insert_at(pos, text):
        insertions[pos] = insertions.get(pos, "") + text

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "delete":
            del_spans.append((a_off[i1], a_off[i2]))
        elif tag == "insert":
            _insert_at(a_off[i1], "".join(b[j1:j2]))
        elif tag == "replace":
            del_spans.append((a_off[i1], a_off[i2]))
            _insert_at(a_off[i2], "".join(b[j1:j2]))
        # 'equal' spans need nothing: their runs stay untouched.
    return del_spans, insertions


def _make_rev_id_allocator(doc):
    """
    Return a callable yielding revision ids unique within the document. Seeds
    above the highest existing w:ins/w:del id so we never collide with revisions
    already in the file.
    """
    max_id = 0
    ins_tag, del_tag, id_attr = qn("w:ins"), qn("w:del"), qn("w:id")
    for el in doc.element.iter():
        if el.tag in (ins_tag, del_tag):
            v = el.get(id_attr)
            if v and v.lstrip("-").isdigit():
                max_id = max(max_id, int(v))
    state = {"n": max_id}

    def _next():
        state["n"] += 1
        return state["n"]

    return _next


def _append_run_text(r, text, deleted=False):
    """
    Append text content to a run element, encoding tabs and newlines as the
    proper <w:tab/>/<w:br/> elements (a literal \\t or \\n inside <w:t> is NOT
    a tab/break to Word). Deleted text is stored as <w:delText> per the schema.
    """
    t_tag = "w:delText" if deleted else "w:t"
    for part in re.split("([\\t\\n])", text):
        if part == "":
            continue
        if part == "\t":
            r.append(OxmlElement("w:tab"))
        elif part == "\n":
            r.append(OxmlElement("w:br"))
        else:
            t = OxmlElement(t_tag)
            t.set(_XML_SPACE, "preserve")  # keep leading/trailing spaces intact
            t.text = part
            r.append(t)


def _set_rev_attrs(el, rev_id, author, date):
    """Stamp the mandatory revision attributes onto a w:ins/w:del element."""
    el.set(qn("w:id"), str(rev_id))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date)


def _normal_run_el(rpr, text):
    """Build a plain <w:r> with the given text, copying run properties (rPr)."""
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    _append_run_text(r, text)
    return r


def _ins_el(rpr, text, rev_id, author, date):
    """Build an <w:ins> wrapping a run that carries `text` (an insertion)."""
    ins = OxmlElement("w:ins")
    _set_rev_attrs(ins, rev_id, author, date)
    ins.append(_normal_run_el(rpr, text))
    return ins


def _del_el(rpr, text, rev_id, author, date):
    """Build a <w:del> wrapping a run whose text is stored as <w:delText>."""
    dele = OxmlElement("w:del")
    _set_rev_attrs(dele, rev_id, author, date)
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    _append_run_text(r, text, deleted=True)
    dele.append(r)
    return dele


def _wrap_run_in_ins(r_el, rev_id, author, date):
    """Wrap an existing run element in <w:ins>, in place."""
    ins = OxmlElement("w:ins")
    _set_rev_attrs(ins, rev_id, author, date)
    r_el.addprevious(ins)
    ins.append(r_el)


def _wrap_run_in_del(r_el, rev_id, author, date):
    """Wrap an existing run element in <w:del> and convert w:t -> w:delText."""
    dele = OxmlElement("w:del")
    _set_rev_attrs(dele, rev_id, author, date)
    r_el.addprevious(dele)
    dele.append(r_el)
    for t in list(r_el.iter(qn("w:t"))):
        dt = OxmlElement("w:delText")
        space = t.get(_XML_SPACE)
        if space:
            dt.set(_XML_SPACE, space)
        dt.text = t.text
        t.getparent().replace(t, dt)


def _para_mark_rpr(p_el):
    """Get or create the paragraph-mark run properties element (w:pPr/w:rPr)."""
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)  # pPr must be the first child of w:p
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        # Schema order inside pPr: rPr comes after every other property,
        # before only sectPr and pPrChange.
        anchor = pPr.find(qn("w:sectPr"))
        if anchor is None:
            anchor = pPr.find(qn("w:pPrChange"))
        if anchor is not None:
            anchor.addprevious(rPr)
        else:
            pPr.append(rPr)
    return rPr


def _mark_paragraph_mark(p_el, kind, rev_id, author, date):
    """
    Mark the paragraph mark (the pilcrow) as a tracked insertion or deletion:
    an empty <w:ins>/<w:del> inside w:pPr/w:rPr. This is how Word records a
    whole-paragraph insert/delete, and it is what makes accept/reject remove
    or restore the paragraph itself rather than just its text.
    `kind` is "w:ins" or "w:del".
    """
    marker = OxmlElement(kind)
    _set_rev_attrs(marker, rev_id, author, date)
    _para_mark_rpr(p_el).insert(0, marker)  # ins/del come first in CT_ParaRPr
    return marker


def _mark_para_boundary(p_el, kind, rev_id, author, date):
    """
    Record the paragraph break belonging to p_el as inserted/deleted, the way
    Word does it. Normally that is p_el's own paragraph mark; but a document's
    FINAL paragraph mark can never be inserted or deleted (Word treats it as
    permanent), so when no paragraph follows, the break that actually
    appears/disappears is the one BEFORE p_el - mark the previous paragraph's
    pilcrow instead. Returns the element whose mark was tagged.
    """
    p_tag = qn("w:p")
    nxt = p_el.getnext()
    if nxt is not None and nxt.tag == p_tag:
        _mark_paragraph_mark(p_el, kind, rev_id, author, date)
        return p_el
    prev = p_el.getprevious()
    if prev is not None and prev.tag == p_tag:
        _mark_paragraph_mark(prev, kind, rev_id, author, date)
        return prev
    # Only paragraph in its parent: no break to mark. Word behaves the same -
    # deleting everything in a one-paragraph document leaves an empty
    # paragraph; the text-level revision alone captures the change.
    return None


def _replace_element(old_el, new_els):
    """Replace old_el in its parent with the ordered list new_els."""
    parent = old_el.getparent()
    idx = parent.index(old_el)
    for j, ne in enumerate(new_els):
        parent.insert(idx + j, ne)
    parent.remove(old_el)


def _apply_tracked_edit(paragraph, del_spans, insertions, author, date, next_id):
    """
    Rewrite a paragraph's top-level runs so that the characters covered by
    `del_spans` (absolute offsets over the concatenated top-level run text)
    become tracked deletions, and each `insertions[pos]` becomes a tracked
    insertion anchored before the character at `pos` (pos == total length
    appends at the end). Characters outside `del_spans` are re-emitted as
    plain runs carrying their ORIGINAL run formatting; runs that need no
    change at all are left completely untouched.
    """
    runs = paragraph.runs
    if not runs:
        # Pure insertion into an empty paragraph.
        for pos in sorted(insertions):
            paragraph._p.append(
                _ins_el(None, insertions[pos], next_id(), author, date))
        return
    texts = [r.text for r in runs]
    starts = []
    acc = 0
    for t in texts:
        starts.append(acc)
        acc += len(t)

    def _deleted(p):
        return any(s <= p < e for s, e in del_spans)

    last = len(runs) - 1
    for i, run in enumerate(runs):
        rs = starts[i]
        text = texts[i]
        re_ = rs + len(text)
        # Insertion anchors this run is responsible for: those falling inside
        # it, plus (last run only) any anchored at/after the very end of text.
        my_ins = {p: insertions[p] for p in insertions
                  if rs <= p < re_ or (i == last and p >= re_)}
        overlaps = any(s < re_ and e > rs for s, e in del_spans)
        if not overlaps and not my_ins:
            continue  # untouched run: original element and formatting survive
        # Split points: run edges, deletion-span edges, insertion anchors.
        bounds = {rs, re_}
        for s, e in del_spans:
            if rs < s < re_:
                bounds.add(s)
            if rs < e < re_:
                bounds.add(e)
        for p in my_ins:
            if rs < p < re_:
                bounds.add(p)
        pts = sorted(bounds)

        r_el = run._r
        rpr = r_el.find(qn("w:rPr"))  # deep-copied inside the builders
        new_els = []
        for a, b in zip(pts, pts[1:]):
            if a in my_ins:
                new_els.append(_ins_el(rpr, my_ins.pop(a), next_id(), author, date))
            seg = text[a - rs: b - rs]
            if not seg:
                continue
            if _deleted(a):
                new_els.append(_del_el(rpr, seg, next_id(), author, date))
            else:
                new_els.append(_normal_run_el(rpr, seg))
        for p in sorted(my_ins):  # anchors at/after the end of the last run
            new_els.append(_ins_el(rpr, my_ins[p], next_id(), author, date))
        _replace_element(r_el, new_els)


def _tracked_replace_in_paragraph(paragraph, find, replace, remaining, author, date, next_id):
    """
    Tracked-changes replace within one paragraph, recorded the way Word
    itself records an edit: `find` and `replace` are DIFFED word-by-word and
    only the words that actually change are wrapped in <w:del>/<w:ins>.
    Unchanged words keep their original runs and formatting. Handles matches
    spanning multiple runs. Returns the number of matches processed.
    """
    if find == replace:
        return 0  # nothing would change; do not record empty revisions
    full = "".join(r.text for r in paragraph.runs)
    matches = []
    start = 0
    while True:
        if remaining is not None and len(matches) >= remaining:
            break
        i = full.find(find, start)
        if i == -1:
            break
        matches.append(i)
        start = i + len(find)
    if not matches:
        return 0
    del_spans, insertions = _word_diff(find, replace)
    # Apply right-to-left so earlier match offsets stay valid: an edit at a
    # later offset never moves characters that sit before it.
    for idx in reversed(matches):
        _apply_tracked_edit(
            paragraph,
            [(idx + s, idx + e) for s, e in del_spans],
            {idx + p: t for p, t in insertions.items()},
            author, date, next_id,
        )
    return len(matches)


def _is_mark_revision(el):
    """True if this w:ins/w:del is a PARAGRAPH-MARK marker (inside pPr/rPr)."""
    parent = el.getparent()
    return (parent is not None and parent.tag == qn("w:rPr")
            and parent.getparent() is not None
            and parent.getparent().tag == qn("w:pPr"))


def _ancestor_paragraph(el):
    """The w:p element containing `el`, or None."""
    p_tag = qn("w:p")
    cur = el.getparent()
    while cur is not None and cur.tag != p_tag:
        cur = cur.getparent()
    return cur


def _revision_text(el):
    """Concatenate the visible text of a revision element (w:t or w:delText)."""
    parts = []
    for t in el.iter(qn("w:t")):
        parts.append(t.text or "")
    for t in el.iter(qn("w:delText")):
        parts.append(t.text or "")
    return "".join(parts)


def _para_plain_text(p_el):
    """All text of a paragraph element, including pending-deleted text."""
    parts = []
    for t in p_el.iter(qn("w:t"), qn("w:delText")):
        parts.append(t.text or "")
    return "".join(parts)


def _body_para_index(doc, p_el):
    """Index of p_el within the document body paragraphs, or None (in a table)."""
    if p_el is None or p_el.getparent() is not doc.element.body:
        return None
    i = 0
    for sib in doc.element.body.iterchildren(qn("w:p")):
        if sib is p_el:
            return i
        i += 1
    return None


def _merge_paragraph_with_next(p_el):
    """
    Remove the paragraph break after p_el: its content flows into the FOLLOWING
    paragraph, whose paragraph mark (and properties) survives - Word's own
    semantics for accepting a deleted paragraph mark / rejecting an inserted
    one. Falls back gracefully when there is no following paragraph.
    """
    pPr_tag, p_tag = qn("w:pPr"), qn("w:p")
    nxt = p_el.getnext()
    if nxt is not None and nxt.tag == p_tag:
        idx = 1 if (len(nxt) and nxt[0].tag == pPr_tag) else 0
        for child in [c for c in p_el if c.tag != pPr_tag]:
            nxt.insert(idx, child)
            idx += 1
        p_el.getparent().remove(p_el)
        return "merged"
    # No following paragraph (end of body / cell): drop the paragraph if it is
    # empty and not the only one, otherwise leave it in place.
    parent = p_el.getparent()
    has_content = any(c.tag != pPr_tag for c in p_el)
    sibling_ps = [c for c in parent if c.tag == p_tag]
    if not has_content and len(sibling_ps) > 1:
        parent.remove(p_el)
        return "removed"
    return "kept"


def _accept_revision(el):
    """Accept one revision element. Returns 'insertion' or 'deletion'."""
    kind = "insertion" if el.tag == qn("w:ins") else "deletion"
    if _is_mark_revision(el):
        p_el = _ancestor_paragraph(el)
        el.getparent().remove(el)
        if kind == "deletion" and p_el is not None:
            _merge_paragraph_with_next(p_el)  # accepted deleted mark: merge
        return kind
    if kind == "insertion":
        parent = el.getparent()
        idx = parent.index(el)
        for child in list(el):          # promote the inserted runs to normal
            parent.insert(idx, child)
            idx += 1
        parent.remove(el)
    else:
        el.getparent().remove(el)       # deleted content disappears
    return kind


def _reject_revision(el):
    """Reject one revision element. Returns 'insertion' or 'deletion'."""
    kind = "insertion" if el.tag == qn("w:ins") else "deletion"
    if _is_mark_revision(el):
        p_el = _ancestor_paragraph(el)
        el.getparent().remove(el)
        if kind == "insertion" and p_el is not None:
            _merge_paragraph_with_next(p_el)  # rejected inserted mark: merge
        return kind
    if kind == "insertion":
        el.getparent().remove(el)       # inserted content disappears
    else:
        parent = el.getparent()
        idx = parent.index(el)
        for child in list(el):          # restore runs, delText -> t
            for dt in child.findall(qn("w:delText")):
                t = OxmlElement("w:t")
                space = dt.get(_XML_SPACE)
                t.set(_XML_SPACE, space if space else "preserve")
                t.text = dt.text
                child.replace(dt, t)
            parent.insert(idx, child)
            idx += 1
        parent.remove(el)
    return kind


def _apply_revisions(doc, accept, ids=None):
    """
    Accept or reject revisions. ids=None means every revision; otherwise only
    those whose w:id is in `ids`. Text-level revisions are processed before
    paragraph-mark revisions so that paragraph merges see final content.
    Returns ({"insertion": n, "deletion": n}, [ids_not_found]).
    """
    wanted = None if ids is None else {str(i) for i in ids}
    found = set()
    content, marks = [], []
    for el in doc.element.iter(qn("w:ins"), qn("w:del")):
        rid = el.get(qn("w:id"))
        if wanted is not None:
            if rid not in wanted:
                continue
            found.add(rid)
        (marks if _is_mark_revision(el) else content).append(el)
    counts = {"insertion": 0, "deletion": 0}
    apply_fn = _accept_revision if accept else _reject_revision
    for el in content + marks:
        counts[apply_fn(el)] += 1
    missing = sorted(wanted - found) if wanted is not None else []
    return counts, missing


# -----------------------------------------------------------------------------
# FINAL-VIEW TEXT RENDERING
#
# python-docx's paragraph.text only sees TOP-LEVEL runs, so any text sitting
# inside a pending <w:ins> is invisible to it (and pending deletions are
# invisible too, because deleted text lives in <w:delText>). Reading a document
# through that lens after tracked edits would make content appear to vanish.
# These helpers render what Word's "No Markup" view shows: the document as it
# would look with every pending change ACCEPTED - insertions included,
# deletions omitted.
# -----------------------------------------------------------------------------
def _final_para_text(p_el):
    """Paragraph text in the final (all-changes-accepted) view."""
    parts = []
    t_tag, tab_tag, br_tag = qn("w:t"), qn("w:tab"), qn("w:br")
    for r in p_el.iter(qn("w:r")):
        for child in r:
            if child.tag == t_tag:
                parts.append(child.text or "")
            elif child.tag == tab_tag:
                parts.append("\t")
            elif child.tag == br_tag:
                parts.append("\n")
            # w:delText is skipped: pending deletion, gone once accepted.
    return "".join(parts)


def _final_cell_text(cell):
    """Table-cell text in the final view (paragraphs joined by newlines)."""
    return "\n".join(_final_para_text(p._p) for p in cell.paragraphs)


def _count_pending_changes(doc):
    """Number of tracked-change elements currently in the document."""
    return sum(1 for _ in doc.element.iter(qn("w:ins"), qn("w:del")))


def _render_linear_text(doc):
    """Full document as plain text, tables flattened to ' | '-joined rows."""
    lines = []
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            lines.append(_final_para_text(block._p))
        elif isinstance(block, Table):
            for row in block.rows:
                lines.append(" | ".join(_final_cell_text(cell) for cell in row.cells))
    return "\n".join(lines)


def _render_structured(doc):
    """
    Ordered block list. Paragraphs carry a para_index into doc.paragraphs and
    tables carry a table_index into doc.tables, so the model can address them
    in later calls (set_paragraph_format, get_tables, etc.).
    """
    blocks = []
    p_i = 0
    t_i = 0
    for b_i, block in enumerate(_iter_block_items(doc)):
        if isinstance(block, Paragraph):
            blocks.append({
                "block_index": b_i,
                "type": "paragraph",
                "para_index": p_i,
                "style": block.style.name if block.style is not None else None,
                "text": _final_para_text(block._p),
            })
            p_i += 1
        elif isinstance(block, Table):
            blocks.append({
                "block_index": b_i,
                "type": "table",
                "table_index": t_i,
                "rows": len(block.rows),
                "cols": len(block.columns),
            })
            t_i += 1
    return blocks


# =============================================================================
# TOOL HANDLERS  (each takes an `args` dict, returns a JSON-serialisable object)
# =============================================================================
def tool_open(args):
    path = _resolve_path(_require(args, "path"))
    if not os.path.isfile(path):
        raise ToolError("File not found: {}".format(path))
    if not path.lower().endswith(".docx"):
        raise ToolError("Only .docx files are supported (got: {})".format(path))
    if len(SESSIONS) >= MAX_SESSIONS:
        raise ToolError(
            "Too many open sessions ({}). Close some with msword_close.".format(
                MAX_SESSIONS
            )
        )
    try:
        doc = docx.Document(path)
    except PackageNotFoundError:
        raise ToolError(
            "Not a valid .docx file (corrupt, or not really Office Open XML)."
        )
    except PermissionError:
        raise ToolError("Permission denied opening: {}".format(path))
    sid = uuid.uuid4().hex[:12]
    SESSIONS[sid] = {"path": path, "doc": doc, "opened_at": _now_iso()}
    log("opened {} as session {}".format(path, sid))
    return {
        "session_id": sid,
        "path": path,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
    }


def tool_close(args):
    sid = _require(args, "session_id")
    if sid not in SESSIONS:
        raise ToolError("Unknown session_id '{}'.".format(sid))
    SESSIONS.pop(sid)
    log("closed session {}".format(sid))
    return {"closed": sid, "note": "Any unsaved changes were discarded."}


def tool_list_sessions(args):
    out = []
    for sid, s in SESSIONS.items():
        out.append({
            "session_id": sid,
            "path": s["path"],
            "opened_at": s["opened_at"],
            "paragraphs": len(s["doc"].paragraphs),
            "tables": len(s["doc"].tables),
        })
    return {"sessions": out, "count": len(out)}


def tool_get_content(args):
    session = _get_session(args)
    doc = session["doc"]
    mode = args.get("mode", "text")
    pending = _count_pending_changes(doc)
    if mode == "text":
        return {"mode": "text", "content": _render_linear_text(doc),
                "pending_changes": pending}
    elif mode == "structured":
        return {"mode": "structured", "blocks": _render_structured(doc),
                "pending_changes": pending}
    else:
        raise ToolError("mode must be 'text' or 'structured' (got '{}')".format(mode))


def tool_search(args):
    session = _get_session(args)
    doc = session["doc"]
    query = _require(args, "query")
    if query == "":
        raise ToolError("query must not be empty")
    case_sensitive = bool(args.get("case_sensitive", False))
    max_results = args.get("max_results")
    if max_results is not None:
        max_results = int(max_results)

    needle = query if case_sensitive else query.lower()

    def _hits(text):
        hay = text if case_sensitive else text.lower()
        positions = []
        start = 0
        while True:
            i = hay.find(needle, start)
            if i == -1:
                break
            positions.append(i)
            start = i + len(needle)
        return positions

    def _snippet(text, i):
        a = max(0, i - SEARCH_CONTEXT_CHARS)
        b = min(len(text), i + len(query) + SEARCH_CONTEXT_CHARS)
        prefix = "..." if a > 0 else ""
        suffix = "..." if b < len(text) else ""
        return prefix + text[a:b] + suffix

    matches = []

    # Body paragraphs (addressable by para_index). Final view: text inside
    # pending insertions is searchable, pending-deleted text is not.
    for p_i, p in enumerate(doc.paragraphs):
        ptext = _final_para_text(p._p)
        for pos in _hits(ptext):
            matches.append({
                "location": "body_paragraph",
                "para_index": p_i,
                "char_offset": pos,
                "snippet": _snippet(ptext, pos),
            })
            if max_results and len(matches) >= max_results:
                return {"query": query, "matches": matches, "truncated": True}

    # Table cells (addressable by table/row/col).
    for t_i, tbl in enumerate(doc.tables):
        for r_i, row in enumerate(tbl.rows):
            for c_i, cell in enumerate(row.cells):
                ctext = _final_cell_text(cell)
                for pos in _hits(ctext):
                    matches.append({
                        "location": "table_cell",
                        "table_index": t_i,
                        "row": r_i,
                        "col": c_i,
                        "char_offset": pos,
                        "snippet": _snippet(ctext, pos),
                    })
                    if max_results and len(matches) >= max_results:
                        return {"query": query, "matches": matches, "truncated": True}

    return {"query": query, "matches": matches, "truncated": False}


def tool_replace_text(args):
    session = _get_session(args)
    doc = session["doc"]
    find = _require(args, "find")
    if find == "":
        raise ToolError("'find' must not be empty")
    replace = args.get("replace", "")
    count = args.get("count")  # max total replacements; None = all
    if count is not None:
        count = int(count)
        if count < 0:
            raise ToolError("'count' must be >= 0")

    track = bool(args.get("track_changes", False))
    author = args.get("author") or AUTHOR
    date = _today_iso()
    next_id = _make_rev_id_allocator(doc) if track else None

    total = 0
    for paragraph in _all_editable_paragraphs(doc):
        remaining = None if count is None else max(0, count - total)
        if remaining == 0:
            break
        if track:
            total += _tracked_replace_in_paragraph(
                paragraph, find, replace, remaining, author, date, next_id
            )
        else:
            total += _replace_in_paragraph(paragraph, find, replace, remaining)

    result = {"find": find, "replace": replace, "replacements": total,
              "tracked": track}
    if track:
        result["author"] = author
        result["date"] = date
    return result


def _has_pending_revisions(p_el):
    """True if the paragraph already contains any tracked change."""
    for _ in p_el.iter(qn("w:ins"), qn("w:del")):
        return True
    return False


def tool_set_paragraph_text(args):
    """
    Replace the entire text of one body paragraph. Tracked mode diffs old vs
    new word-by-word, so only the words that actually change become revisions.
    """
    session = _get_session(args)
    doc = session["doc"]
    para_index = int(_require(args, "para_index"))
    if not 0 <= para_index < len(doc.paragraphs):
        raise ToolError(
            "para_index out of range (0..{}).".format(len(doc.paragraphs) - 1)
        )
    new_text = _require(args, "text")
    if not isinstance(new_text, str):
        raise ToolError("'text' must be a string")
    paragraph = doc.paragraphs[para_index]
    track = bool(args.get("track_changes", False))

    if not track:
        # Silent rewrite: drop all content children (runs, hyperlinks, any
        # pending revisions), keep the paragraph properties/style.
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        if new_text:
            paragraph.add_run(new_text)
        return {"para_index": para_index, "tracked": False, "text": new_text}

    warning = None
    if _has_pending_revisions(paragraph._p):
        warning = ("Paragraph already contains pending tracked changes; the "
                   "diff was computed against its unrevised text only. "
                   "Accept/reject existing changes first for a clean result.")
    old_text = "".join(r.text for r in paragraph.runs)
    if old_text == new_text:
        return {"para_index": para_index, "tracked": True, "changed": False}
    author = args.get("author") or AUTHOR
    date = _today_iso()
    next_id = _make_rev_id_allocator(doc)
    del_spans, insertions = _word_diff(old_text, new_text)
    _apply_tracked_edit(paragraph, del_spans, insertions, author, date, next_id)
    result = {"para_index": para_index, "tracked": True, "changed": True,
              "author": author, "date": date,
              "deleted_spans": len(del_spans), "inserted_spans": len(insertions)}
    if warning:
        result["warning"] = warning
    return result


def tool_insert_paragraph(args):
    """
    Insert a new paragraph before/after a body paragraph (or append at the
    end). Tracked mode marks the text AND the paragraph mark as inserted, so
    rejecting removes the whole paragraph - Word's native behaviour.
    """
    session = _get_session(args)
    doc = session["doc"]
    text = args.get("text", "")
    style = args.get("style")
    position = args.get("position", "after")
    if position not in ("before", "after"):
        raise ToolError("position must be 'before' or 'after'")
    para_index = args.get("para_index")

    try:
        if para_index is None:
            new_p = doc.add_paragraph(text, style=style)
        else:
            para_index = int(para_index)
            if not 0 <= para_index < len(doc.paragraphs):
                raise ToolError(
                    "para_index out of range (0..{}).".format(len(doc.paragraphs) - 1)
                )
            ref = doc.paragraphs[para_index]
            new_p = ref.insert_paragraph_before(text, style)
            if position == "after":
                ref._p.addnext(new_p._p)  # relocate: after the anchor instead
    except KeyError:
        raise ToolError("Unknown paragraph style: '{}'".format(style))

    track = bool(args.get("track_changes", False))
    author = args.get("author") or AUTHOR
    date = _today_iso()
    if track:
        next_id = _make_rev_id_allocator(doc)
        for r_el in [r._r for r in new_p.runs]:
            _wrap_run_in_ins(r_el, next_id(), author, date)
        _mark_para_boundary(new_p._p, "w:ins", next_id(), author, date)

    new_index = next(
        i for i, p in enumerate(doc.paragraphs) if p._p is new_p._p
    )
    result = {"para_index": new_index, "text": text, "tracked": track}
    if track:
        result["author"] = author
        result["date"] = date
    return result


def tool_delete_paragraph(args):
    """
    Delete a body paragraph. Tracked mode marks the text AND the paragraph
    mark as deleted (shown struck out; accepting removes the paragraph and
    merges with the next, rejecting restores it) - Word's native behaviour.
    """
    session = _get_session(args)
    doc = session["doc"]
    para_index = int(_require(args, "para_index"))
    if not 0 <= para_index < len(doc.paragraphs):
        raise ToolError(
            "para_index out of range (0..{}).".format(len(doc.paragraphs) - 1)
        )
    paragraph = doc.paragraphs[para_index]
    preview = paragraph.text
    track = bool(args.get("track_changes", False))

    if track:
        author = args.get("author") or AUTHOR
        date = _today_iso()
        next_id = _make_rev_id_allocator(doc)
        for r_el in [r._r for r in paragraph.runs]:
            _wrap_run_in_del(r_el, next_id(), author, date)
        _mark_para_boundary(paragraph._p, "w:del", next_id(), author, date)
        return {"para_index": para_index, "tracked": True, "author": author,
                "date": date, "text": preview}

    body = doc.element.body
    if (paragraph._p.getparent() is body
            and len(body.findall(qn("w:p"))) <= 1):
        raise ToolError(
            "Refusing to delete the document's only paragraph "
            "(a body should keep at least one). Clear its text instead."
        )
    paragraph._p.getparent().remove(paragraph._p)
    return {"para_index": para_index, "tracked": False, "text": preview}


def tool_add_paragraph(args):
    session = _get_session(args)
    doc = session["doc"]
    text = args.get("text", "")
    style = args.get("style")
    try:
        if style:
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(text)
    except KeyError:
        raise ToolError("Unknown paragraph style: '{}'".format(style))
    return {"para_index": len(doc.paragraphs) - 1, "text": text}


def tool_add_heading(args):
    session = _get_session(args)
    doc = session["doc"]
    text = args.get("text", "")
    level = int(args.get("level", 1))
    if not 0 <= level <= 9:
        raise ToolError("heading level must be between 0 (Title) and 9")
    try:
        doc.add_heading(text, level=level)
    except KeyError:
        raise ToolError(
            "Heading style for level {} is not defined in this document.".format(level)
        )
    return {"para_index": len(doc.paragraphs) - 1, "level": level, "text": text}


def tool_add_table(args):
    session = _get_session(args)
    doc = session["doc"]
    data = args.get("data")
    style = args.get("style")  # e.g. "Table Grid"

    if data is not None:
        if not isinstance(data, list) or not data or not all(
            isinstance(row, list) for row in data
        ):
            raise ToolError("'data' must be a non-empty list of row lists")
        rows = len(data)
        cols = max(len(r) for r in data)
    else:
        rows = int(_require(args, "rows"))
        cols = int(_require(args, "cols"))
        if rows < 1 or cols < 1:
            raise ToolError("rows and cols must both be >= 1")

    table = doc.add_table(rows=rows, cols=cols)
    if style:
        try:
            table.style = style
        except KeyError:
            raise ToolError("Unknown table style: '{}'".format(style))
    if data is not None:
        for r_i, row in enumerate(data):
            for c_i in range(cols):
                value = row[c_i] if c_i < len(row) else ""
                table.rows[r_i].cells[c_i].text = "" if value is None else str(value)

    return {"table_index": len(doc.tables) - 1, "rows": rows, "cols": cols}


def tool_get_tables(args):
    session = _get_session(args)
    doc = session["doc"]
    table_index = args.get("table_index")

    def _dump(tbl):
        return [[cell.text for cell in row.cells] for row in tbl.rows]

    if table_index is not None:
        table_index = int(table_index)
        if not 0 <= table_index < len(doc.tables):
            raise ToolError("table_index out of range (0..{})".format(len(doc.tables) - 1))
        return {"table_index": table_index, "cells": _dump(doc.tables[table_index])}

    return {"tables": [
        {"table_index": i, "cells": _dump(t)} for i, t in enumerate(doc.tables)
    ]}


# Accept both US and Australian spellings of "centre".
_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "centre": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def tool_set_paragraph_format(args):
    session = _get_session(args)
    doc = session["doc"]
    para_index = int(_require(args, "para_index"))
    if not 0 <= para_index < len(doc.paragraphs):
        raise ToolError(
            "para_index out of range (0..{}).".format(len(doc.paragraphs) - 1)
        )
    paragraph = doc.paragraphs[para_index]
    applied = {}

    # Character formatting is applied to every run in the paragraph.
    for attr in ("bold", "italic", "underline"):
        if attr in args and args[attr] is not None:
            value = bool(args[attr])
            for run in paragraph.runs:
                setattr(run, attr, value)
            applied[attr] = value

    if "alignment" in args and args["alignment"] is not None:
        key = str(args["alignment"]).lower()
        if key not in _ALIGN_MAP:
            raise ToolError(
                "alignment must be one of: {}".format(", ".join(sorted(_ALIGN_MAP)))
            )
        paragraph.alignment = _ALIGN_MAP[key]
        applied["alignment"] = key

    if "style" in args and args["style"] is not None:
        try:
            paragraph.style = args["style"]
        except KeyError:
            raise ToolError("Unknown paragraph style: '{}'".format(args["style"]))
        applied["style"] = args["style"]

    if not applied:
        raise ToolError(
            "No formatting supplied. Provide one or more of: "
            "bold, italic, underline, alignment, style."
        )
    return {"para_index": para_index, "applied": applied}


def tool_save(args):
    session = _get_session(args)
    doc = session["doc"]
    save_as = args.get("path")
    if save_as:
        target = _resolve_path(save_as)
        if not target.lower().endswith(".docx"):
            raise ToolError("Save path must end in .docx")
        parent = os.path.dirname(target)
        if parent and not os.path.isdir(parent):
            raise ToolError("Destination folder does not exist: {}".format(parent))
    else:
        target = session["path"]

    try:
        doc.save(target)
    except PermissionError:
        raise ToolError(
            "Permission denied saving to {} (is it open in Word?).".format(target)
        )
    except OSError as e:
        raise ToolError("Could not save to {}: {}".format(target, e))

    if save_as:
        session["path"] = target  # rebind session to the new file
    log("saved session {} -> {}".format(args.get("session_id"), target))
    return {"saved": target}


def tool_list_changes(args):
    session = _get_session(args)
    doc = session["doc"]
    changes = []
    for el in doc.element.iter(qn("w:ins"), qn("w:del")):
        mark = _is_mark_revision(el)
        p_el = _ancestor_paragraph(el)
        para_text = _para_plain_text(p_el) if p_el is not None else ""
        changes.append({
            "type": "insertion" if el.tag == qn("w:ins") else "deletion",
            # 'text' = words inserted/deleted; 'paragraph_mark' = the pilcrow
            # itself (a whole-paragraph insert/delete carries one of each).
            "scope": "paragraph_mark" if mark else "text",
            "id": el.get(qn("w:id")),
            "author": el.get(qn("w:author")),
            "date": el.get(qn("w:date")),
            "text": "" if mark else _revision_text(el),
            "para_index": _body_para_index(doc, p_el),  # None = inside a table
            "paragraph_text": (para_text[:117] + "...") if len(para_text) > 120
                              else para_text,
        })
    return {"changes": changes, "count": len(changes)}


def _ids_arg(args):
    ids = _require(args, "ids")
    if not isinstance(ids, list) or not ids:
        raise ToolError("'ids' must be a non-empty array of change ids")
    return ids


def tool_accept_changes(args):
    session = _get_session(args)
    counts, missing = _apply_revisions(session["doc"], accept=True,
                                       ids=_ids_arg(args))
    result = {"accepted_insertions": counts["insertion"],
              "accepted_deletions": counts["deletion"]}
    if missing:
        result["ids_not_found"] = missing
    return result


def tool_reject_changes(args):
    session = _get_session(args)
    counts, missing = _apply_revisions(session["doc"], accept=False,
                                       ids=_ids_arg(args))
    result = {"rejected_insertions": counts["insertion"],
              "rejected_deletions": counts["deletion"]}
    if missing:
        result["ids_not_found"] = missing
    return result


def tool_accept_all_changes(args):
    session = _get_session(args)
    counts, _ = _apply_revisions(session["doc"], accept=True)
    return {"accepted_insertions": counts["insertion"],
            "accepted_deletions": counts["deletion"]}


def tool_reject_all_changes(args):
    session = _get_session(args)
    counts, _ = _apply_revisions(session["doc"], accept=False)
    return {"rejected_insertions": counts["insertion"],
            "rejected_deletions": counts["deletion"]}


# =============================================================================
# TOOL REGISTRY  (name -> (handler, JSON-Schema inputSchema, description))
# =============================================================================
TOOLS = [
    {
        "name": "msword_open",
        "description": "Open a .docx file into an in-memory session and return a session_id used by all other tools.",
        "handler": tool_open,
        "inputSchema": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "Absolute or ~ path to a .docx file."}
            },
            "required": ["path"],
        },
    },
    {
        "name": "msword_close",
        "description": "Close a session and free its document. Unsaved changes are discarded.",
        "handler": tool_close,
        "inputSchema": {
            "type": "object",
            "properties": {"session_id": {"type": "string"}},
            "required": ["session_id"],
        },
    },
    {
        "name": "msword_list_sessions",
        "description": "List all currently open document sessions.",
        "handler": tool_list_sessions,
        "inputSchema": {"type": "object", "properties": {}},
    },
    {
        "name": "msword_get_content",
        "description": "Read a document. mode='text' returns linear plain text (tables flattened); mode='structured' returns an ordered block list with para_index/table_index for addressing. Text is the FINAL view (pending tracked insertions shown, pending deletions hidden - like Word's 'No Markup'); the result's pending_changes count says whether revisions are outstanding.",
        "handler": tool_get_content,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "mode": {"type": "string", "enum": ["text", "structured"], "default": "text"},
            },
            "required": ["session_id"],
        },
    },
    {
        "name": "msword_search",
        "description": "Find text in body paragraphs and table cells. Returns match locations and context snippets.",
        "handler": tool_search,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "query": {"type": "string"},
                "case_sensitive": {"type": "boolean", "default": False},
                "max_results": {"type": "integer", "description": "Optional cap on results."},
            },
            "required": ["session_id", "query"],
        },
    },
    {
        "name": "msword_replace_text",
        "description": "Replace text across body paragraphs and table cells. Handles matches spanning multiple runs. Set track_changes=true to record the edit as Word tracked changes: find/replace are diffed word-by-word so only the words that actually change are marked as deleted/inserted (like editing in Word with Track Changes on) - safe to pass whole sentences or paragraphs. An empty 'replace' with track_changes=true is a tracked deletion.",
        "handler": tool_replace_text,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "find": {"type": "string"},
                "replace": {"type": "string", "default": ""},
                "count": {"type": "integer", "description": "Optional max total replacements; omit to replace all."},
                "track_changes": {"type": "boolean", "default": False, "description": "Record edits as Word tracked changes."},
                "author": {"type": "string", "description": "Override the tracked-change author for this call (defaults to the server's --author)."},
            },
            "required": ["session_id", "find"],
        },
    },
    {
        "name": "msword_set_paragraph_text",
        "description": "Replace the entire text of one body paragraph (para_index from get_content structured). With track_changes=true the old and new text are diffed word-by-word and only the changed words become tracked deletions/insertions - use this to 'rewrite this paragraph' as a clean Word revision. Untracked, it silently replaces the paragraph's content (keeping its style).",
        "handler": tool_set_paragraph_text,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "para_index": {"type": "integer"},
                "text": {"type": "string", "description": "The paragraph's full new text ('' deletes all its text)."},
                "track_changes": {"type": "boolean", "default": False},
                "author": {"type": "string", "description": "Override the tracked-change author for this call."},
            },
            "required": ["session_id", "para_index", "text"],
        },
    },
    {
        "name": "msword_insert_paragraph",
        "description": "Insert a new paragraph before/after a body paragraph (omit para_index to append at the end). With track_changes=true the text AND the paragraph mark are recorded as a Word tracked insertion, so rejecting the change removes the whole paragraph.",
        "handler": tool_insert_paragraph,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "text": {"type": "string", "default": ""},
                "para_index": {"type": "integer", "description": "Anchor paragraph; omit to append at the document end."},
                "position": {"type": "string", "enum": ["before", "after"], "default": "after"},
                "style": {"type": "string", "description": "Optional paragraph style, e.g. 'Normal', 'Heading 2'."},
                "track_changes": {"type": "boolean", "default": False},
                "author": {"type": "string", "description": "Override the tracked-change author for this call."},
            },
            "required": ["session_id"],
        },
    },
    {
        "name": "msword_delete_paragraph",
        "description": "Delete a body paragraph. With track_changes=true the text AND the paragraph mark are recorded as a Word tracked deletion (text shows struck out; accepting removes the paragraph entirely, rejecting restores it). Untracked, the paragraph is removed immediately.",
        "handler": tool_delete_paragraph,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "para_index": {"type": "integer"},
                "track_changes": {"type": "boolean", "default": False},
                "author": {"type": "string", "description": "Override the tracked-change author for this call."},
            },
            "required": ["session_id", "para_index"],
        },
    },
    {
        "name": "msword_add_paragraph",
        "description": "Append a paragraph to the end of the document. Optional paragraph style name.",
        "handler": tool_add_paragraph,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "text": {"type": "string", "default": ""},
                "style": {"type": "string", "description": "Optional style, e.g. 'Normal', 'Quote'."},
            },
            "required": ["session_id"],
        },
    },
    {
        "name": "msword_add_heading",
        "description": "Append a heading. level 0 = Title, 1..9 = Heading 1..9.",
        "handler": tool_add_heading,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "text": {"type": "string", "default": ""},
                "level": {"type": "integer", "minimum": 0, "maximum": 9, "default": 1},
            },
            "required": ["session_id"],
        },
    },
    {
        "name": "msword_add_table",
        "description": "Append a table. Either supply 'data' (list of row lists) to fill it, or 'rows' and 'cols' for an empty grid. Optional table style e.g. 'Table Grid'.",
        "handler": tool_add_table,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "data": {
                    "type": "array",
                    "items": {"type": "array", "items": {"type": "string"}},
                    "description": "Row-major cell values.",
                },
                "rows": {"type": "integer", "minimum": 1},
                "cols": {"type": "integer", "minimum": 1},
                "style": {"type": "string"},
            },
            "required": ["session_id"],
        },
    },
    {
        "name": "msword_get_tables",
        "description": "Return table contents as arrays of cell text. Omit table_index for all tables.",
        "handler": tool_get_tables,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "table_index": {"type": "integer"},
            },
            "required": ["session_id"],
        },
    },
    {
        "name": "msword_set_paragraph_format",
        "description": "Apply simple formatting to a body paragraph (addressed by para_index from get_content structured). Any of: bold, italic, underline, alignment (left/center/centre/right/justify), style.",
        "handler": tool_set_paragraph_format,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "para_index": {"type": "integer"},
                "bold": {"type": "boolean"},
                "italic": {"type": "boolean"},
                "underline": {"type": "boolean"},
                "alignment": {"type": "string", "enum": ["left", "center", "centre", "right", "justify"]},
                "style": {"type": "string"},
            },
            "required": ["session_id", "para_index"],
        },
    },
    {
        "name": "msword_save",
        "description": "Save the session's document. Omit 'path' to save in place, or supply a .docx 'path' to save-as (the session then tracks the new file).",
        "handler": tool_save,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "path": {"type": "string", "description": "Optional save-as path ending in .docx."},
            },
            "required": ["session_id"],
        },
    },
    {
        "name": "msword_list_changes",
        "description": "List every tracked change in the document with id, author, date, changed text, scope ('text', or 'paragraph_mark' for an inserted/deleted pilcrow - a whole-paragraph insert/delete has one of each), para_index (None = inside a table) and surrounding paragraph text. The ids feed msword_accept_changes / msword_reject_changes.",
        "handler": tool_list_changes,
        "inputSchema": {
            "type": "object",
            "properties": {"session_id": {"type": "string"}},
            "required": ["session_id"],
        },
    },
    {
        "name": "msword_accept_changes",
        "description": "Accept specific tracked changes by id (from msword_list_changes): inserted text becomes normal text, deleted text is removed. Accepting a deleted paragraph mark merges the paragraph with the following one, as Word does.",
        "handler": tool_accept_changes,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "ids": {"type": "array", "items": {"type": "string"},
                        "description": "Change ids to accept."},
            },
            "required": ["session_id", "ids"],
        },
    },
    {
        "name": "msword_reject_changes",
        "description": "Reject specific tracked changes by id (from msword_list_changes): inserted text is removed, deleted text is restored. Rejecting an inserted paragraph mark removes that paragraph break, as Word does.",
        "handler": tool_reject_changes,
        "inputSchema": {
            "type": "object",
            "properties": {
                "session_id": {"type": "string"},
                "ids": {"type": "array", "items": {"type": "string"},
                        "description": "Change ids to reject."},
            },
            "required": ["session_id", "ids"],
        },
    },
    {
        "name": "msword_accept_all_changes",
        "description": "Accept every tracked change: inserted text/paragraphs are kept as normal content and deleted text/paragraphs are removed.",
        "handler": tool_accept_all_changes,
        "inputSchema": {
            "type": "object",
            "properties": {"session_id": {"type": "string"}},
            "required": ["session_id"],
        },
    },
    {
        "name": "msword_reject_all_changes",
        "description": "Reject every tracked change: inserted text/paragraphs are removed and deleted text/paragraphs are restored.",
        "handler": tool_reject_all_changes,
        "inputSchema": {
            "type": "object",
            "properties": {"session_id": {"type": "string"}},
            "required": ["session_id"],
        },
    },
]

_TOOL_BY_NAME = {t["name"]: t for t in TOOLS}


# =============================================================================
# JSON-RPC / MCP DISPATCH
# =============================================================================
def _jsonrpc_result(msg_id, result):
    return {"jsonrpc": "2.0", "id": msg_id, "result": result}


def _jsonrpc_error(msg_id, code, message):
    return {"jsonrpc": "2.0", "id": msg_id, "error": {"code": code, "message": message}}


def handle_message(msg):
    """
    Process one decoded JSON-RPC message. Returns a response dict, or None for
    notifications (which must not be answered).
    """
    method = msg.get("method")
    msg_id = msg.get("id")
    is_notification = "id" not in msg
    params = msg.get("params") or {}

    if method == "initialize":
        client_proto = params.get("protocolVersion") or PROTOCOL_VERSION_FALLBACK
        return _jsonrpc_result(msg_id, {
            "protocolVersion": client_proto,
            "capabilities": {"tools": {}},
            "serverInfo": {"name": SERVER_NAME, "version": SERVER_VERSION},
        })

    if method in ("notifications/initialized", "initialized"):
        return None  # notification, no response

    if method == "ping":
        return _jsonrpc_result(msg_id, {})

    if method == "tools/list":
        listed = [
            {"name": t["name"], "description": t["description"], "inputSchema": t["inputSchema"]}
            for t in TOOLS
        ]
        return _jsonrpc_result(msg_id, {"tools": listed})

    if method == "tools/call":
        name = params.get("name")
        arguments = params.get("arguments") or {}
        tool = _TOOL_BY_NAME.get(name)
        if tool is None:
            return _jsonrpc_result(msg_id, {
                "content": [{"type": "text", "text": "Unknown tool: {}".format(name)}],
                "isError": True,
            })
        try:
            result = tool["handler"](arguments)
            text = json.dumps(result, indent=2, ensure_ascii=False)
            return _jsonrpc_result(msg_id, {
                "content": [{"type": "text", "text": text}],
                "isError": False,
            })
        except ToolError as e:
            return _jsonrpc_result(msg_id, {
                "content": [{"type": "text", "text": "Error: {}".format(e)}],
                "isError": True,
            })
        except Exception as e:  # unexpected - log full trace to stderr, return summary
            log("UNEXPECTED in tool {}:\n{}".format(name, traceback.format_exc()))
            return _jsonrpc_result(msg_id, {
                "content": [{"type": "text", "text": "Internal error: {}".format(e)}],
                "isError": True,
            })

    # Unknown method.
    if is_notification:
        return None
    return _jsonrpc_error(msg_id, -32601, "Method not found: {}".format(method))


def serve():
    """Main stdio loop: read newline-delimited JSON-RPC, dispatch, reply."""
    docx_ver, lxml_ver = _versions()
    log("starting")
    log("interpreter: {}".format(sys.executable))
    log("python-docx {} / lxml {}".format(docx_ver, lxml_ver))
    log("tracked-change author: {}".format(AUTHOR))
    if DOCUMENT_ROOT:
        log("path sandbox DOCUMENT_ROOT = {}".format(DOCUMENT_ROOT))

    for line in sys.stdin:
        line = line.strip()
        if not line:
            continue
        try:
            msg = json.loads(line)
        except json.JSONDecodeError:
            # Cannot know the id; emit a parse error with null id per spec.
            sys.stdout.write(json.dumps(_jsonrpc_error(None, -32700, "Parse error")) + "\n")
            sys.stdout.flush()
            continue

        try:
            response = handle_message(msg)
        except Exception as e:  # last-resort guard so the server never dies
            log("FATAL in handle_message:\n{}".format(traceback.format_exc()))
            response = _jsonrpc_error(msg.get("id"), -32603, "Internal error: {}".format(e))

        if response is not None:
            sys.stdout.write(json.dumps(response, ensure_ascii=False) + "\n")
            sys.stdout.flush()

    log("stdin closed, exiting")


# =============================================================================
# SELF-TEST  (--check): offline round-trip so a single transfer can be validated
# =============================================================================
def run_check():
    import tempfile
    global DOCUMENT_ROOT
    docx_ver, lxml_ver = _versions()
    print("[check] interpreter : {}".format(sys.executable))
    print("[check] python-docx : {}".format(docx_ver))
    print("[check] lxml        : {}".format(lxml_ver))

    tmpdir = tempfile.mkdtemp(prefix="msword_check_")
    # The self-test sandboxes itself to its own temp folder so it can run
    # before the endpoint's real DOCUMENT_ROOT exists.
    DOCUMENT_ROOT = tmpdir
    print("[check] sandbox     : {} (self-test only)".format(tmpdir))
    path = os.path.join(tmpdir, "roundtrip.docx")
    ok = True
    try:
        # Build a document exercising each editing path.
        d = docx.Document()
        d.add_heading("Title Here", level=0)
        d.add_paragraph("The quick brown fox jumps over the lazy dog.")
        d.add_table(rows=2, cols=2)
        d.tables[0].rows[0].cells[0].text = "old value"
        d.save(path)

        # Reopen through the tool layer.
        opened = tool_open({"path": path})
        sid = opened["session_id"]
        assert opened["paragraphs"] >= 2, "expected >=2 paragraphs"
        assert opened["tables"] == 1, "expected 1 table"

        # Search should find the fox sentence.
        s = tool_search({"session_id": sid, "query": "brown fox"})
        assert s["matches"], "search found nothing"

        # Replace across body + table cell.
        r1 = tool_replace_text({"session_id": sid, "find": "quick", "replace": "slow"})
        assert r1["replacements"] == 1, "body replace count wrong"
        r2 = tool_replace_text({"session_id": sid, "find": "old value", "replace": "new value"})
        assert r2["replacements"] == 1, "table-cell replace count wrong"

        # Append + format.
        tool_add_paragraph({"session_id": sid, "text": "Appended paragraph."})
        pi = tool_get_content({"session_id": sid, "mode": "structured"})["blocks"]
        last_para = max(b["para_index"] for b in pi if b["type"] == "paragraph")
        tool_set_paragraph_format({"session_id": sid, "para_index": last_para,
                                   "bold": True, "alignment": "centre"})

        # Save-as and reopen to confirm persistence of the edits.
        path2 = os.path.join(tmpdir, "roundtrip2.docx")
        tool_save({"session_id": sid, "path": path2})
        tool_close({"session_id": sid})

        reopened = tool_open({"path": path2})
        text = tool_get_content({"session_id": reopened["session_id"], "mode": "text"})["content"]
        assert "slow brown fox" in text, "body edit did not persist"
        assert "new value" in text, "table edit did not persist"
        assert "Appended paragraph." in text, "append did not persist"
        tool_close({"session_id": reopened["session_id"]})

        print("[check] round-trip: PASS")

        # --- Tracked changes 1: simple word replace, accept + reject --------
        tpath = os.path.join(tmpdir, "tracked.docx")
        td = docx.Document()
        td.add_paragraph("The report is DRAFT and confidential.")
        td.save(tpath)

        ts = tool_open({"path": tpath})
        tsid = ts["session_id"]
        rt = tool_replace_text({"session_id": tsid, "find": "DRAFT",
                                "replace": "FINAL", "track_changes": True,
                                "author": "Self Test"})
        assert rt["replacements"] == 1 and rt["tracked"], "tracked replace failed"

        listed = tool_list_changes({"session_id": tsid})
        kinds = sorted(c["type"] for c in listed["changes"])
        assert kinds == ["deletion", "insertion"], \
            "expected one insertion + one deletion, got {}".format(kinds)
        assert all(c["author"] == "Self Test" for c in listed["changes"]), \
            "author not stamped"
        assert all(c["scope"] == "text" for c in listed["changes"]), \
            "unexpected paragraph-mark change"
        assert all(c["para_index"] == 0 for c in listed["changes"]), \
            "para_index wrong in list_changes"

        tpath2 = os.path.join(tmpdir, "tracked2.docx")
        tool_save({"session_id": tsid, "path": tpath2})
        tool_close({"session_id": tsid})

        # Accept path: FINAL stays, DRAFT gone.
        a = tool_open({"path": tpath2})
        tool_accept_all_changes({"session_id": a["session_id"]})
        atext = tool_get_content({"session_id": a["session_id"], "mode": "text"})["content"]
        assert "FINAL" in atext and "DRAFT" not in atext, "accept-all wrong"
        assert tool_list_changes({"session_id": a["session_id"]})["count"] == 0, \
            "changes left after accept-all"
        tool_close({"session_id": a["session_id"]})

        # Reject path: DRAFT restored, FINAL gone.
        r = tool_open({"path": tpath2})
        tool_reject_all_changes({"session_id": r["session_id"]})
        rtext = tool_get_content({"session_id": r["session_id"], "mode": "text"})["content"]
        assert "DRAFT" in rtext and "FINAL" not in rtext, "reject-all wrong"
        tool_close({"session_id": r["session_id"]})

        print("[check] tracked-changes basic: PASS")

        # --- Tracked changes 2: word-level diff on a whole-sentence replace -
        # The review record must contain ONLY the changed words, never the
        # whole sentence as one deletion + one insertion.
        old_s = "The quick brown fox jumps over the lazy dog."
        new_s = "The quick red fox leaps over the lazy dog."
        d2path = os.path.join(tmpdir, "diff.docx")
        d2 = docx.Document()
        d2.add_paragraph(old_s)
        d2.save(d2path)
        s2 = tool_open({"path": d2path})["session_id"]
        rt = tool_replace_text({"session_id": s2, "find": old_s, "replace": new_s,
                                "track_changes": True, "author": "Self Test"})
        assert rt["replacements"] == 1, "sentence replace failed"
        ch = tool_list_changes({"session_id": s2})["changes"]
        dels = sorted(c["text"] for c in ch if c["type"] == "deletion")
        inss = sorted(c["text"] for c in ch if c["type"] == "insertion")
        assert dels == ["brown", "jumps"], \
            "diff not word-level, deletions: {}".format(dels)
        assert inss == ["leaps", "red"], \
            "diff not word-level, insertions: {}".format(inss)
        tool_accept_all_changes({"session_id": s2})
        t2 = tool_get_content({"session_id": s2, "mode": "text"})["content"]
        assert new_s in t2, "accepted diff text wrong"
        tool_close({"session_id": s2})

        # Pure insertion inside a match must terminate and land correctly.
        d3path = os.path.join(tmpdir, "insonly.docx")
        d3 = docx.Document()
        d3.add_paragraph("Feed the lazy dog today.")
        d3.save(d3path)
        s3 = tool_open({"path": d3path})["session_id"]
        rt = tool_replace_text({"session_id": s3, "find": "lazy dog",
                                "replace": "lazy old dog",
                                "track_changes": True})
        assert rt["replacements"] == 1, "insert-only replace failed"
        tool_accept_all_changes({"session_id": s3})
        t3 = tool_get_content({"session_id": s3, "mode": "text"})["content"]
        assert "Feed the lazy old dog today." in t3, "insert-only result wrong"
        tool_close({"session_id": s3})

        # Formatting on unchanged words must survive (multi-run match).
        d4path = os.path.join(tmpdir, "fmt.docx")
        d4 = docx.Document()
        p4 = d4.add_paragraph("")
        r4 = p4.add_run("Alpha ")
        r4.bold = True
        p4.add_run("beta gamma")
        d4.save(d4path)
        s4 = tool_open({"path": d4path})["session_id"]
        tool_replace_text({"session_id": s4, "find": "Alpha beta",
                           "replace": "Alpha delta", "track_changes": True})
        doc4 = SESSIONS[s4]["doc"]
        first_run = doc4.paragraphs[0].runs[0]
        assert first_run.text == "Alpha " and first_run.bold, \
            "unchanged bold run was disturbed by the diff"
        tool_accept_all_changes({"session_id": s4})
        t4 = tool_get_content({"session_id": s4, "mode": "text"})["content"]
        assert "Alpha delta gamma" in t4, "multi-run diff accept wrong"
        tool_close({"session_id": s4})

        print("[check] tracked-changes word-diff: PASS")

        # --- Tracked changes 3: set_paragraph_text (rewrite a paragraph) ----
        d5path = os.path.join(tmpdir, "setpara.docx")
        d5 = docx.Document()
        d5.add_paragraph("One two three four")
        d5.save(d5path)
        s5 = tool_open({"path": d5path})["session_id"]
        tool_set_paragraph_text({"session_id": s5, "para_index": 0,
                                 "text": "One 2 three four five",
                                 "track_changes": True})
        ch = tool_list_changes({"session_id": s5})["changes"]
        dels = sorted(c["text"] for c in ch if c["type"] == "deletion")
        inss = sorted(c["text"] for c in ch if c["type"] == "insertion")
        assert dels == ["two"], "set_paragraph_text deletions: {}".format(dels)
        assert inss == [" five", "2"], \
            "set_paragraph_text insertions: {}".format(inss)
        tool_accept_all_changes({"session_id": s5})
        t5 = tool_get_content({"session_id": s5, "mode": "text"})["content"]
        assert "One 2 three four five" in t5, "set_paragraph_text accept wrong"
        tool_close({"session_id": s5})

        # --- Tracked changes 4: paragraph insert/delete incl. the mark ------
        def _three_para_doc(name):
            pth = os.path.join(tmpdir, name)
            dd = docx.Document()
            dd.add_paragraph("One")
            dd.add_paragraph("Three")
            dd.save(pth)
            return pth

        # Inserted paragraph: reject-all must remove it entirely.
        s6 = tool_open({"path": _three_para_doc("pins.docx")})["session_id"]
        ins_res = tool_insert_paragraph({"session_id": s6, "text": "Two",
                                         "para_index": 0, "position": "after",
                                         "track_changes": True})
        assert ins_res["para_index"] == 1, "inserted paragraph at wrong index"
        doc6 = SESSIONS[s6]["doc"]
        # Final view: pending-inserted text must be visible when reading.
        assert [_final_para_text(p._p) for p in doc6.paragraphs] == \
            ["One", "Two", "Three"], "inserted paragraph not in final view"
        ch = tool_list_changes({"session_id": s6})["changes"]
        assert any(c["scope"] == "paragraph_mark" and c["type"] == "insertion"
                   for c in ch), "paragraph mark not tracked as inserted"
        p6path = os.path.join(tmpdir, "pins2.docx")
        tool_save({"session_id": s6, "path": p6path})
        tool_close({"session_id": s6})

        s6r = tool_open({"path": p6path})["session_id"]
        tool_reject_all_changes({"session_id": s6r})
        doc6r = SESSIONS[s6r]["doc"]
        assert [p.text for p in doc6r.paragraphs] == ["One", "Three"], \
            "rejected inserted paragraph did not disappear"
        tool_close({"session_id": s6r})

        s6a = tool_open({"path": p6path})["session_id"]
        tool_accept_all_changes({"session_id": s6a})
        doc6a = SESSIONS[s6a]["doc"]
        assert [p.text for p in doc6a.paragraphs] == ["One", "Two", "Three"], \
            "accepted inserted paragraph wrong"
        assert tool_list_changes({"session_id": s6a})["count"] == 0
        tool_close({"session_id": s6a})

        # Deleted paragraph: accept-all must remove it entirely.
        p7path = os.path.join(tmpdir, "pdel.docx")
        d7 = docx.Document()
        d7.add_paragraph("One")
        d7.add_paragraph("Two")
        d7.add_paragraph("Three")
        d7.save(p7path)
        s7 = tool_open({"path": p7path})["session_id"]
        tool_delete_paragraph({"session_id": s7, "para_index": 1,
                               "track_changes": True})
        ch = tool_list_changes({"session_id": s7})["changes"]
        assert any(c["scope"] == "paragraph_mark" and c["type"] == "deletion"
                   for c in ch), "paragraph mark not tracked as deleted"
        assert any(c["scope"] == "text" and c["text"] == "Two"
                   for c in ch), "paragraph text not tracked as deleted"
        p7path2 = os.path.join(tmpdir, "pdel2.docx")
        tool_save({"session_id": s7, "path": p7path2})
        tool_close({"session_id": s7})

        s7a = tool_open({"path": p7path2})["session_id"]
        tool_accept_all_changes({"session_id": s7a})
        doc7a = SESSIONS[s7a]["doc"]
        assert [p.text for p in doc7a.paragraphs] == ["One", "Three"], \
            "accepted deleted paragraph did not disappear"
        tool_close({"session_id": s7a})

        s7r = tool_open({"path": p7path2})["session_id"]
        tool_reject_all_changes({"session_id": s7r})
        doc7r = SESSIONS[s7r]["doc"]
        assert [p.text for p in doc7r.paragraphs] == ["One", "Two", "Three"], \
            "rejected deleted paragraph not restored"
        tool_close({"session_id": s7r})

        # End-of-document edges: the FINAL paragraph mark is permanent in
        # Word, so deleting the last paragraph / appending a paragraph must
        # mark the PREVIOUS paragraph's pilcrow instead.
        p9path = os.path.join(tmpdir, "plast.docx")
        d9 = docx.Document()
        d9.add_paragraph("One")
        d9.add_paragraph("Two")
        d9.save(p9path)
        s9 = tool_open({"path": p9path})["session_id"]
        tool_delete_paragraph({"session_id": s9, "para_index": 1,
                               "track_changes": True})
        p9path2 = os.path.join(tmpdir, "plast2.docx")
        tool_save({"session_id": s9, "path": p9path2})
        tool_close({"session_id": s9})
        s9a = tool_open({"path": p9path2})["session_id"]
        tool_accept_all_changes({"session_id": s9a})
        assert [p.text for p in SESSIONS[s9a]["doc"].paragraphs] == ["One"], \
            "accepted delete-of-last-paragraph wrong"
        tool_close({"session_id": s9a})
        s9r = tool_open({"path": p9path2})["session_id"]
        tool_reject_all_changes({"session_id": s9r})
        assert [p.text for p in SESSIONS[s9r]["doc"].paragraphs] == ["One", "Two"], \
            "rejected delete-of-last-paragraph wrong"
        tool_close({"session_id": s9r})

        s10 = tool_open({"path": p9path})["session_id"]
        tool_insert_paragraph({"session_id": s10, "text": "Appended",
                               "track_changes": True})
        p10path = os.path.join(tmpdir, "pappend.docx")
        tool_save({"session_id": s10, "path": p10path})
        tool_close({"session_id": s10})
        s10r = tool_open({"path": p10path})["session_id"]
        tool_reject_all_changes({"session_id": s10r})
        assert [p.text for p in SESSIONS[s10r]["doc"].paragraphs] == ["One", "Two"], \
            "rejected appended paragraph wrong"
        tool_close({"session_id": s10r})
        s10a = tool_open({"path": p10path})["session_id"]
        tool_accept_all_changes({"session_id": s10a})
        assert [p.text for p in SESSIONS[s10a]["doc"].paragraphs] == \
            ["One", "Two", "Appended"], "accepted appended paragraph wrong"
        tool_close({"session_id": s10a})

        print("[check] tracked-changes paragraphs: PASS")

        # --- Tracked changes 5: accept/reject individual changes by id ------
        p8path = os.path.join(tmpdir, "byid.docx")
        d8 = docx.Document()
        d8.add_paragraph("Alpha Bravo")
        d8.add_paragraph("Charlie Delta")
        d8.save(p8path)
        s8 = tool_open({"path": p8path})["session_id"]
        tool_replace_text({"session_id": s8, "find": "Bravo",
                           "replace": "Bravissimo", "track_changes": True})
        tool_replace_text({"session_id": s8, "find": "Charlie",
                           "replace": "Charles", "track_changes": True})
        ch = tool_list_changes({"session_id": s8})["changes"]
        ids_p0 = [c["id"] for c in ch if c["para_index"] == 0]
        ids_p1 = [c["id"] for c in ch if c["para_index"] == 1]
        assert ids_p0 and ids_p1, "per-paragraph ids not found"
        tool_accept_changes({"session_id": s8, "ids": ids_p0})
        tool_reject_changes({"session_id": s8, "ids": ids_p1})
        t8 = tool_get_content({"session_id": s8, "mode": "text"})["content"]
        assert "Alpha Bravissimo" in t8, "accept-by-id wrong"
        assert "Charlie Delta" in t8 and "Charles" not in t8, "reject-by-id wrong"
        assert tool_list_changes({"session_id": s8})["count"] == 0, \
            "changes left after per-id accept/reject"
        tool_close({"session_id": s8})

        print("[check] tracked-changes accept/reject by id: PASS")
    except Exception as e:
        ok = False
        print("[check] round-trip: FAIL -> {}".format(e))
        traceback.print_exc()
    finally:
        # Clean up temp files; leave nothing behind.
        try:
            for f in os.listdir(tmpdir):
                os.remove(os.path.join(tmpdir, f))
            os.rmdir(tmpdir)
        except Exception:
            pass

    if ok:
        print("[check] ALL CHECKS PASSED")
        return 0
    print("[check] CHECKS FAILED")
    return 1


def main():
    parser = argparse.ArgumentParser(
        description="MS Word (.docx) python-docx MCP stdio server. "
                    "With no arguments it runs as an MCP server on stdin/stdout."
    )
    parser.add_argument(
        "--check", action="store_true",
        help="Run an offline open/edit/save/reopen self-test and exit."
    )
    parser.add_argument(
        "--author", default=os.environ.get("MSWORD_AUTHOR"), metavar="NAME",
        help="Author name stamped on tracked changes (falls back to the "
             "MSWORD_AUTHOR environment variable, then the "
             "TRACKED_CHANGE_AUTHOR config value)."
    )
    parser.add_argument(
        "--document-root", default=os.environ.get("MSWORD_DOCUMENT_ROOT"),
        metavar="DIR",
        help="REQUIRED path sandbox (unless the DOCUMENT_ROOT config constant "
             "is set, or via the MSWORD_DOCUMENT_ROOT environment variable): "
             "the server refuses to open or save any file outside this "
             "directory tree, and refuses to start without one. The model "
             "chooses open/save paths, so an unconfined server could "
             "read/write any .docx this account can."
    )
    args = parser.parse_args()

    global AUTHOR, DOCUMENT_ROOT
    if args.author:
        AUTHOR = args.author
    if args.document_root:
        DOCUMENT_ROOT = args.document_root

    if args.check:
        sys.exit(run_check())

    # File access is confined to DOCUMENT_ROOT, so a root is mandatory.
    if not DOCUMENT_ROOT:
        log("FATAL: no document root configured. Pass --document-root, set the "
            "MSWORD_DOCUMENT_ROOT environment variable, or set the "
            "DOCUMENT_ROOT constant in this file. The server only opens/saves "
            ".docx files inside that folder and will not start without one.")
        sys.exit(2)
    if not os.path.isdir(DOCUMENT_ROOT):
        log("FATAL: the configured document root does not exist or is not a "
            "directory: {}".format(DOCUMENT_ROOT))
        sys.exit(2)

    try:
        serve()
    except KeyboardInterrupt:
        log("interrupted, exiting")


if __name__ == "__main__":
    main()
